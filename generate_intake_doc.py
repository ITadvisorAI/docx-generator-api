from flask import Flask, request, jsonify, send_from_directory, abort
from docx import Document
import os
import shutil
from threading import Thread
import re
import requests
import traceback
import json

print("🚀 Flask is starting...")

app = Flask(__name__)

# ✅ Sanitize session_id for folder paths
def sanitize_session_id(session_id):
    return re.sub(r"[^\w\-]", "_", session_id)

# ✅ Safe filename from email
def clean_email_for_filename(email):
    return email.replace("@", "_at_").replace(".", "_dot_")

# ✅ Background thread for document generation
def process_intake_document(data, file_name, output_path):
    try:
        print("🚀 Thread started: processing intake document")

        session_id = data.get('session_id')
        intake = data.get('intake_answers', {})
        files = data.get('files', [])

        safe_session_id = sanitize_session_id(session_id)
        folder_path = os.path.join("temp_sessions", f"Temp_{safe_session_id}")
        os.makedirs(folder_path, exist_ok=True)
        print(f"📁 Folder created or exists: {folder_path}")

        # Save debug snapshot
        try:
            with open(os.path.join(folder_path, "debug.json"), "w") as f:
                json.dump(data, f, indent=2)
            print("📝 Saved debug input JSON")
        except Exception as debug_error:
            print("⚠️ Failed to save debug input:", debug_error)

        # Load and copy template
        script_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(script_dir, "intakeform.docx")
        print(f"🔍 Looking for template at: {template_path}")
        if not os.path.exists(template_path):
            print(f"❌ Template file missing: {template_path}")
            return

        shutil.copy(template_path, output_path)
        print(f"📄 Template copied to: {output_path}")

        doc = Document(output_path)

        doc.add_heading("Selected Programs", level=1)
        selected_categories = intake.get("selected_categories", [])
        selected_programs = intake.get("selected_programs", {})

        for category in selected_categories:
            doc.add_paragraph(category, style="ListBullet")
            for program in selected_programs.get(category, []):
                doc.add_paragraph(f"  - {program}", style="ListBullet2")

        doc.add_heading("Transformation Questions", level=1)
        doc.add_paragraph(f"1. {intake.get('q1', '')}")
        doc.add_paragraph(f"2. {intake.get('q2', '')}")
        doc.add_paragraph(f"3. {intake.get('q3', '')}")
        doc.add_paragraph(f"4. {intake.get('q4', '')}")
        doc.add_paragraph(f"5. {intake.get('q5', '')}")

        if files:
            doc.add_heading("Uploaded Files", level=1)
            for f in files:
                doc.add_paragraph(f"{f.get('name', 'Unknown')} ({f.get('type', '')})", style="ListBullet")
                doc.add_paragraph(f"URL: {f.get('url', '')}", style="Normal")

        doc.save(output_path)
        print(f"✅ DOCX successfully saved: {output_path}")

    except Exception as e:
        print("❌ Exception in process_intake_document:")
        traceback.print_exc()

# ✅ Intake DOCX Generator Route
@app.route('/generate_intake', methods=['POST'])
def generate_intake():
    try:
        print("📥 /generate_intake endpoint STARTED")

        raw = request.data.decode("utf-8")
        print("📦 RAW REQUEST:")
        print(raw)

        data = request.get_json(force=True)
        print("✅ Parsed JSON:")
        print(data)

        session_id = data.get('session_id')
        email = data.get('email')
        intake = data.get('intake_answers', {})

        if not session_id or not email or not intake:
            print("❌ Missing required fields in request.")
            return jsonify({
                "error": "Missing session_id, email, or intake_answers"
            }), 400

        safe_session_id = sanitize_session_id(session_id)
        email_id = clean_email_for_filename(email)
        file_name = f"intake_{sanitize_session_id(session_id)}_{email_id}.docx"
        file_url = f"https://docx-generator-api.onrender.com/files/Temp_{safe_session_id}/{file_name}"
        output_path = os.path.join("temp_sessions", f"Temp_{safe_session_id}", file_name)

        print(f"🧵 Launching background thread for file: {file_name}")
        Thread(target=process_intake_document, args=(data, file_name, output_path)).start()

        return jsonify({
            "status": "processing",
            "session_id": session_id,
            "file_name": file_name,
            "file_url": file_url
        }), 202

    except Exception as e:
        print("❌ Exception in /generate_intake:")
        traceback.print_exc()
        return jsonify({"error": "Internal server error", "details": str(e)}), 500

# ✅ Route to serve generated DOCX files
@app.route('/files/<path:filename>', methods=['GET'])
def serve_generated_file(filename):
    try:
        directory = os.path.join(os.getcwd(), 'temp_sessions')
        full_path = os.path.join(directory, filename)

        if not os.path.isfile(full_path):
            print(f"❌ File not found: {full_path}")
            abort(404)

        print(f"📤 Serving file: {full_path}")
        return send_from_directory(directory, filename, as_attachment=False)

    except Exception as e:
        print(f"❌ Exception in /files route:")
        traceback.print_exc()
        abort(500)

# ✅ Proxy route to forward email to Make.com webhook
@app.route('/start_session', methods=['POST'])
def start_session():
    try:
        data = request.get_json(force=True)
        print("📩 Received session start request:")
        print(data)

        make_webhook = 'https://hook.us2.make.com/1ivi9q9x6l253tikb557hemgtl7n2bv9'
        r = requests.post(make_webhook, json=data)

        if r.status_code == 200:
            print("✅ Session initiated via Make.com")
            return jsonify({"message": "Session started"}), 200
        else:
            print(f"❌ Failed to POST to Make.com: {r.status_code}")
            return jsonify({"error": "Make webhook failed", "status": r.status_code}), r.status_code

    except Exception as e:
        print("❌ Exception in /start_session:")
        traceback.print_exc()
        return jsonify({"error": "Internal server error", "details": str(e)}), 500

# ✅ Health check endpoint
@app.route('/healthz', methods=['GET'])
def health_check():
    return "ok", 200

# ✅ Run app (Render will assign dynamic port via $PORT env var)
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
