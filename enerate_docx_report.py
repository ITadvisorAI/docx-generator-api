from docx import Document
from docx.shared import Inches
import os

def generate_docx_report_with_charts(session_id):
    # Paths
    session_path = f"temp_sessions/{session_id}"
    chart_path = os.path.join(session_path, "charts")
    output_path = os.path.join(session_path, f"IT_Infrastructure_Current_Status_Report_{session_id}.docx")
    template_path = "templates/IT_Infrastructure_Report_Template.docx"

    # Ensure output directory exists
    os.makedirs(session_path, exist_ok=True)

    # Load the template or create a new document
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
        doc.add_heading("IT Infrastructure Assessment Report", level=1)

    # Add chart section
    doc.add_heading("📊 Infrastructure Visualizations", level=2)

    # List of expected chart files
    chart_files = [
        ("HW Tier Distribution", "hw_tier_distribution.png"),
        ("HW Environment Distribution", "hw_environment_distribution.png"),
        ("HW Device Type vs Tier", "hw_device_type_vs_tier.png"),
        ("SW Tier Distribution", "sw_tier_distribution.png"),
        ("SW Environment Distribution", "sw_environment_distribution.png")
    ]

    for title, filename in chart_files:
        chart_file = os.path.join(chart_path, filename)
        if os.path.exists(chart_file):
            doc.add_heading(title, level=3)
            doc.add_picture(chart_file, width=Inches(5.5))
        else:
            doc.add_paragraph(f"⚠️ {title} chart not found.")

    # Save the final report
    doc.save(output_path)
    return output_path
